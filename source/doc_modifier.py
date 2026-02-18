
import logging
import json
from docx import Document

from generate_contents import ResumeFormattedResponse, JobDetails, JobBlock

from typing import Any, Dict, List

class DocModifier:

    def __init__(self, doc_path: str) -> None:
        self.logger = logging.getLogger(f"uvicorn.{__name__}")
        self.logger.info(f"Initializing DocModifier with doc path: {doc_path}")
        self._doc = Document(doc_path)
        self.sections = self._get_sections()

        self.blocks = []
        for name, section in self.sections.items():
            item = JobBlock(name, [i.text for i in section])
            self.blocks.append(item)

    def _get_sections(self) -> Dict[str, Any]:
        sections = {}
        i = 1
        while i < len(self._doc.paragraphs) - 1:
            prev_paragraph = self._doc.paragraphs[i - 1]
            this_paragraph = self._doc.paragraphs[i]
            next_paragraph = self._doc.paragraphs[i + 1]
            if this_paragraph.style.name == "Normal" and next_paragraph.style.name == "List Paragraph":
                company = prev_paragraph.text.strip().split(",")[0]
                sections[company] = []
                
                j = 1
                while next_paragraph.style.name == "List Paragraph":
                    sections[company].append(next_paragraph)
                    next_paragraph = self._doc.paragraphs[i + 1 + j]
                    j += 1
                i += j
                
            i += 1
        return sections

    def modify_section(self, section: Dict[str, Any], changes: JobBlock):
        self.logger.info(f"Modifying section {changes.company} with changes: {changes.bullet_points}")
        if len(self.sections[section]) != len(changes.bullet_points):
            # TODO Handle this
            return

        for i, bullet_point in enumerate(self.sections[section]):
            bullet_point.text = changes.bullet_points[i]

    def modify_sections(self, new_sections: List[JobBlock]):
        if len(self.sections) != len(new_sections):
            raise ValueError(f"Number of sections in resume ({len(self.sections)}) does not match number of sections in changes ({len(new_sections)})")
        
        for old_section, new_section in zip(self.sections, new_sections):
            self.modify_section(old_section, new_section)
    
    def save(self, path: str):
        self._doc.save(path)

if __name__ == "__main__":
    modifier = DocModifier("Isaac_Trussell.docx")
    print(modifier.sections)